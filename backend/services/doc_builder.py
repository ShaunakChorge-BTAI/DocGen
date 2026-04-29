import io
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .template_loader import load_brand_config

# ── Shared brand constants ─────────────────────────────────────────────────────
COLOR_H1 = RGBColor(0x1A, 0x3C, 0x5E)
COLOR_H2 = RGBColor(0x2E, 0x7D, 0xB2)
COLOR_H3 = RGBColor(0x1A, 0x3C, 0x5E)
COLOR_TABLE_HEADER_BG = "E8F4FD"
FONT_NAME = "Calibri"

PDF_PRIMARY   = "#1A3C5E"
PDF_SECONDARY = "#2E7DB2"
PDF_ACCENT    = "#E8F4FD"


_DOC_TYPE_TITLES = {
    "BRD": "Business Requirements Document",
    "FSD": "Functional Specification Document",
    "SRS": "Software Requirements Specification",
    "User Manual": "User Manual",
    "Product Brochure": "Product Brochure",
}


def _substitute_placeholders(markdown: str, version: str, brand_config: dict, doc_type: str = "") -> str:
    """Replace LLM-output placeholders with real values from brand config and current date."""
    today = datetime.utcnow().strftime("%d-%b-%Y")
    company = brand_config.get("company_name", "")
    product = brand_config.get("product_name", "")
    author = brand_config.get("default_author", "Implementation Team")
    author_role = brand_config.get("default_author_role", "Business Analyst")
    reviewer = brand_config.get("default_reviewer", "TBD")
    approver = brand_config.get("default_approver", "TBD")
    full_title = _DOC_TYPE_TITLES.get(doc_type, doc_type)
    doc_title = f"{full_title} — {product}" if product else full_title

    subs = {
        "[DOCUMENT TITLE]": doc_title,
        "[DATE]": today,
        "[VERSION]": version,
        "[COMPANY NAME]": company,
        "[COMPANY]": company,
        "[PROJECT NAME]": product,
        "[PRODUCT NAME]": product,
        "[AUTHOR NAME]": author,
        "[AUTHOR]": author,
        "[ROLE]": author_role,
        "[REVIEWER NAME]": reviewer,
        "[APPROVER NAME]": approver,
    }
    for placeholder, value in subs.items():
        markdown = markdown.replace(placeholder, value)
    return markdown


# ══════════════════════════════════════════════════════════════════════════════
# Title-page helpers
# ══════════════════════════════════════════════════════════════════════════════

def _extract_header_table(markdown: str) -> tuple[dict, str]:
    """
    Pull the Document Header table out of the markdown.
    Returns (header_data_dict, remaining_markdown_without_title_and_header_section).
    """
    lines = markdown.splitlines()
    header_data: dict = {}
    result_lines: list[str] = []
    title_skipped = False
    in_header_section = False
    i = 0
    while i < len(lines):
        line = lines[i]
        # Drop the first H1 (document title line)
        if not title_skipped and line.startswith("# "):
            title_skipped = True
            i += 1
            continue
        # Detect "Document Header" heading
        if re.match(r"^#{1,3}\s+Document Header", line, re.IGNORECASE):
            in_header_section = True
            i += 1
            continue
        if in_header_section:
            if line.startswith("|"):
                cells = [c.strip() for c in line.strip().strip("|").split("|")]
                is_separator = all(re.match(r"^[-:]+$", c) for c in cells if c)
                if not is_separator and len(cells) >= 2:
                    field, value = cells[0], cells[1]
                    if field.lower() not in ("field", ""):
                        header_data[field] = value
                i += 1
                continue
            elif line.startswith("---") or line.strip() == "":
                i += 1
                continue
            else:
                in_header_section = False
                result_lines.append(line)
        else:
            result_lines.append(line)
        i += 1
    return header_data, "\n".join(result_lines)


def _build_title_page(doc: Document, doc_type: str, version: str, brand_config: dict, header_data: dict):
    """Render a professional, branded title page."""
    company_name = brand_config.get("company_name", "")
    product_name = brand_config.get("product_name", "")
    full_title    = _DOC_TYPE_TITLES.get(doc_type, doc_type)
    today         = datetime.utcnow().strftime("%d-%b-%Y")
    doc_title     = header_data.get("Document Title") or f"{full_title} — {product_name}"
    author        = brand_config.get("default_author", "")
    author_role   = brand_config.get("default_author_role", "")
    prepared_by   = header_data.get("Prepared By") or (f"{author}  ·  {author_role}" if author_role else author)

    section = doc.sections[0]
    avail_w = section.page_width - section.left_margin - section.right_margin

    # ── 1. TOP BANNER — company name (navy) + product (blue) ──
    banner = doc.add_table(rows=2, cols=1)
    banner.style = "Table Grid"
    for row in banner.rows:
        for cell in row.cells:
            _remove_cell_borders(cell)

    c1 = banner.rows[0].cells[0]
    _set_cell_bg(c1, "1A3C5E")
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.space_before = Pt(10)
    p1.paragraph_format.space_after  = Pt(10)
    p1.paragraph_format.left_indent  = Cm(0.5)
    r1 = p1.add_run(company_name.upper())
    r1.font.name = FONT_NAME; r1.font.size = Pt(11)
    r1.font.bold = True; r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    c2 = banner.rows[1].cells[0]
    _set_cell_bg(c2, "2E7DB2")
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(5)
    p2.paragraph_format.space_after  = Pt(5)
    p2.paragraph_format.left_indent  = Cm(0.5)
    r2 = p2.add_run(product_name)
    r2.font.name = FONT_NAME; r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0xE8, 0xF4, 0xFD)

    # ── 2. SPACER ──────────────────────────────────────────────
    _add_spacer_para(doc, 28)

    # ── 3. DOCUMENT TYPE LABEL ─────────────────────────────────
    lbl = doc.add_paragraph()
    lbl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lbl.paragraph_format.left_indent = Cm(0.5)
    lbl.paragraph_format.space_after = Pt(4)
    lr = lbl.add_run(doc_type)
    lr.font.name = FONT_NAME; lr.font.size = Pt(12)
    lr.font.color.rgb = COLOR_H2

    # ── 4. BIG DOCUMENT TITLE ──────────────────────────────────
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tp.paragraph_format.left_indent  = Cm(0.5)
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after  = Pt(10)
    tr = tp.add_run(full_title)
    tr.font.name = FONT_NAME; tr.font.size = Pt(26)
    tr.font.bold = True; tr.font.color.rgb = COLOR_H1

    # ── 5. THICK ACCENT RULE ───────────────────────────────────
    rule = _add_rule(doc, "2E7DB2", 3)
    rule.paragraph_format.left_indent = Cm(0.5)

    # ── 6. PRODUCT + VERSION / DATE SUBTITLE ──────────────────
    _add_spacer_para(doc, 10)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub.paragraph_format.left_indent = Cm(0.5)
    sub.paragraph_format.space_after = Pt(3)
    sr = sub.add_run(product_name)
    sr.font.name = FONT_NAME; sr.font.size = Pt(14)
    sr.font.color.rgb = COLOR_H2

    meta_line = doc.add_paragraph()
    meta_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_line.paragraph_format.left_indent = Cm(0.5)
    meta_line.paragraph_format.space_after = Pt(40)
    mr = meta_line.add_run(
        f"{header_data.get('Version', version)}  ·  {header_data.get('Date', today)}"
    )
    mr.font.name = FONT_NAME; mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # ── 7. DOCUMENT DETAILS TABLE ──────────────────────────────
    rows_data = [
        ("Document Title",  doc_title),
        ("Version",         header_data.get("Version",  version)),
        ("Date",            header_data.get("Date",     today)),
        ("Prepared By",     prepared_by),
        ("Reviewed By",     header_data.get("Reviewed By",  brand_config.get("default_reviewer", "TBD"))),
        ("Approved By",     header_data.get("Approved By",  brand_config.get("default_approver", "TBD"))),
        ("Classification",  header_data.get("Classification", "INTERNAL — CONFIDENTIAL")),
    ]

    meta = doc.add_table(rows=len(rows_data) + 1, cols=2)
    meta.style = "Table Grid"

    # Header row spanning both columns
    hcell = meta.rows[0].cells[0].merge(meta.rows[0].cells[1])
    _set_cell_bg(hcell, "1A3C5E")
    hp = hcell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.left_indent  = Cm(0.3)
    hp.paragraph_format.space_before = Pt(7)
    hp.paragraph_format.space_after  = Pt(7)
    hrun = hp.add_run("Document Details")
    hrun.font.name = FONT_NAME; hrun.font.size = Pt(10)
    hrun.font.bold = True; hrun.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, (field, value) in enumerate(rows_data):
        row = meta.rows[idx + 1]
        bg  = "EBF5FB" if idx % 2 == 0 else "FFFFFF"

        fc = row.cells[0]
        _set_cell_bg(fc, bg)
        fp = fc.paragraphs[0]
        fp.paragraph_format.left_indent  = Cm(0.3)
        fp.paragraph_format.space_before = Pt(6)
        fp.paragraph_format.space_after  = Pt(6)
        frun = fp.add_run(field)
        frun.font.name = FONT_NAME; frun.font.size = Pt(10)
        frun.font.bold = True; frun.font.color.rgb = COLOR_H1

        vc = row.cells[1]
        _set_cell_bg(vc, bg)
        vp = vc.paragraphs[0]
        vp.paragraph_format.left_indent  = Cm(0.3)
        vp.paragraph_format.space_before = Pt(6)
        vp.paragraph_format.space_after  = Pt(6)
        vrun = vp.add_run(value)
        vrun.font.name = FONT_NAME; vrun.font.size = Pt(10)
        vrun.font.color.rgb = RGBColor(0x2D, 0x3D, 0x4E)

    # Fix column widths: label 35%, value 65%
    label_w = int(avail_w * 0.35)
    value_w = avail_w - label_w
    for row in meta.rows:
        row.cells[0].width = label_w
        row.cells[1].width = value_w

    # ── 8. CLASSIFICATION BADGE ────────────────────────────────
    _add_spacer_para(doc, 20)
    badge_tbl = doc.add_table(rows=1, cols=1)
    badge_tbl.style = "Table Grid"
    bc = badge_tbl.rows[0].cells[0]
    _set_cell_bg(bc, "EBF5FB")
    _remove_cell_borders(bc)
    bp = bc.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.paragraph_format.space_before = Pt(7)
    bp.paragraph_format.space_after  = Pt(7)
    br = bp.add_run("INTERNAL  —  CONFIDENTIAL")
    br.font.name = FONT_NAME; br.font.size = Pt(9)
    br.font.bold = True; br.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # ── 9. PAGE BREAK ──────────────────────────────────────────
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# DOCX builder
# ══════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for bname in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{bname}")
        b.set(qn("w:val"), "none")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _add_spacer_para(doc: Document, space_after_pt: float = 8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after_pt)
    return p


def _add_rule(doc: Document, color_hex: str = "2E7DB2", thickness_pt: float = 2):
    """Horizontal rule via paragraph bottom border."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(thickness_pt * 4)))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def _set_run_font(run, size_pt: int, bold=False, color: RGBColor | None = None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _add_heading(doc: Document, text: str, level: int):
    style_map = {
        1: ("Heading 1", 16, COLOR_H1),
        2: ("Heading 2", 13, COLOR_H2),
        3: ("Heading 3", 11, COLOR_H3),
        4: ("Heading 4", 10, COLOR_H3),
    }
    style_name, size, color = style_map.get(level, ("Normal", 11, None))
    para = doc.add_paragraph(style=style_name)
    run = para.add_run(text)
    _set_run_font(run, size, bold=True, color=color)
    return para


def _parse_inline(para, text: str):
    # Match bold before italic to avoid ** being consumed by * patterns
    pattern = r"(\*\*[^*]+\*\*|\*[^*\n]+\*|`[^`\n]+`)"
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = Pt(11)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
            run.font.name = FONT_NAME
            run.font.size = Pt(11)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            run = para.add_run(part)
            run.font.name = FONT_NAME
            run.font.size = Pt(11)


def _add_table(doc: Document, lines: list[str]):
    data_rows = [l for l in lines if not re.match(r"^\|[\s\-:|]+\|$", l)]
    if not data_rows:
        return

    def parse_row(line: str) -> list[str]:
        return [c.strip() for c in line.strip().strip("|").split("|")]

    rows = [parse_row(r) for r in data_rows]
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= col_count:
                break
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            if row_idx == 0:
                _set_cell_bg(cell, COLOR_TABLE_HEADER_BG)
                run = para.add_run(cell_text)
                _set_run_font(run, 11, bold=True, color=COLOR_H1)
            else:
                run = para.add_run(cell_text)
                _set_run_font(run, 11)

    doc.add_paragraph()


def _add_header_footer(doc: Document, doc_type: str, version: str = "v1.0", company_name: str = ""):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    htable = header.add_table(
        1, 3,
        width=section.page_width - section.left_margin - section.right_margin,
    )
    htable.style = "Table Grid"
    left_cell = htable.rows[0].cells[0]
    right_cell = htable.rows[0].cells[2]

    left_run = left_cell.paragraphs[0].add_run(doc_type)
    _set_run_font(left_run, 9, bold=True, color=COLOR_H1)

    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_run = right_para.add_run(version)
    _set_run_font(right_run, 9, color=COLOR_H2)

    for row in htable.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for border_name in ("top", "left", "bottom", "right"):
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "none")
                tcBorders.append(border)
            tcPr.append(tcBorders)

    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    company_run = para.add_run(f"{company_name} | {doc_type} | {version} | Page ")
    _set_run_font(company_run, 8)

    for field, code in [("begin", "PAGE"), ("end", "")]:
        if field == "begin":
            fldChar = OxmlElement("w:fldChar")
            fldChar.set(qn("w:fldCharType"), "begin")
            instrText = OxmlElement("w:instrText")
            instrText.text = "PAGE"
            fldChar2 = OxmlElement("w:fldChar")
            fldChar2.set(qn("w:fldCharType"), "end")
            run_el = para.add_run()
            run_el._r.append(fldChar)
            run_el._r.append(instrText)
            run_el._r.append(fldChar2)
            _set_run_font(run_el, 8)

    of_run = para.add_run(" of ")
    _set_run_font(of_run, 8)

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "begin")
    instrText2 = OxmlElement("w:instrText")
    instrText2.text = "NUMPAGES"
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run_el2 = para.add_run()
    run_el2._r.append(fldChar3)
    run_el2._r.append(instrText2)
    run_el2._r.append(fldChar4)
    _set_run_font(run_el2, 8)

    conf_para = footer.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = conf_para.add_run(
        "This document is confidential and intended for authorized personnel only."
    )
    _set_run_font(conf_run, 7, color=COLOR_H2)


def build_docx(markdown_content: str, doc_type: str, version: str = "v1.0") -> io.BytesIO:
    brand_config = load_brand_config()
    markdown_content = _substitute_placeholders(markdown_content, version, brand_config, doc_type)
    company_name = brand_config.get("company_name", "")

    doc = Document()
    section = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(2.5))

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing = Pt(11 * 1.15)
    pf.space_after = Pt(6)

    _add_header_footer(doc, doc_type, version, company_name)

    # Extract Document Header metadata and strip it from the markdown body
    header_data, body_markdown = _extract_header_table(markdown_content)
    _build_title_page(doc, doc_type, version, brand_config, header_data)

    lines = body_markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("#### "):
            _add_heading(doc, line[5:].strip(), 4)
        elif line.startswith("### "):
            _add_heading(doc, line[4:].strip(), 3)
        elif line.startswith("## "):
            _add_heading(doc, line[3:].strip(), 2)
        elif line.startswith("# "):
            _add_heading(doc, line[2:].strip(), 1)
        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue
        elif re.match(r"^-{3,}$", line.strip()):
            pass
        elif line.startswith("> "):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)
            run = para.add_run(line[2:].strip())
            run.italic = True
            run.font.name = FONT_NAME
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_H2
        elif re.match(r"^  [-*] |^\t[-*] ", line):
            para = doc.add_paragraph(style="List Bullet 2")
            _parse_inline(para, re.sub(r"^[\s\t]+[-*] ", "", line))
        elif re.match(r"^  \d+\. |^\t\d+\. ", line):
            para = doc.add_paragraph(style="List Number 2")
            _parse_inline(para, re.sub(r"^[\s\t]+\d+\. ", "", line))
        elif line.startswith("- ") or line.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            _parse_inline(para, line[2:].strip())
        elif re.match(r"^\d+\. ", line):
            para = doc.add_paragraph(style="List Number")
            _parse_inline(para, re.sub(r"^\d+\. ", "", line).strip())
        elif line.strip():
            para = doc.add_paragraph()
            _parse_inline(para, line.strip())
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
# PDF builder (Phase 3 — reportlab)
# ══════════════════════════════════════════════════════════════════════════════

def build_pdf(markdown_content: str, doc_type: str, version: str = "v1.0") -> io.BytesIO:
    """Converts markdown to a branded PDF using reportlab."""
    brand_config = load_brand_config()
    markdown_content = _substitute_placeholders(markdown_content, version, brand_config, doc_type)
    company_name = brand_config.get("company_name", "")
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor, white
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, ListFlowable, ListItem,
    )

    primary   = HexColor(PDF_PRIMARY)
    secondary = HexColor(PDF_SECONDARY)
    accent    = HexColor(PDF_ACCENT)

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        topMargin=3 * cm,
        bottomMargin=2.5 * cm,
        title=f"{doc_type} {version}",
    )

    ss = getSampleStyleSheet()

    def style(name, **kw):
        return ParagraphStyle(name, parent=ss["Normal"], **kw)

    h1_style  = style("H1",  fontSize=16, leading=20, textColor=primary,   spaceAfter=10, fontName="Helvetica-Bold")
    h2_style  = style("H2",  fontSize=13, leading=17, textColor=secondary,  spaceAfter=8,  fontName="Helvetica-Bold")
    h3_style  = style("H3",  fontSize=11, leading=14, textColor=primary,   spaceAfter=6,  fontName="Helvetica-Bold")
    body_style = style("Body", fontSize=11, leading=14, spaceAfter=6)
    bullet_style = style("Bullet", fontSize=11, leading=14, leftIndent=16, spaceAfter=3)
    footer_style = style("Footer", fontSize=8, textColor=secondary, alignment=TA_CENTER)
    header_style = style("Header", fontSize=9, textColor=primary, fontName="Helvetica-Bold")

    def _header_footer(canvas, doc):
        canvas.saveState()
        # Header bar
        canvas.setFillColor(accent)
        canvas.rect(2 * cm, A4[1] - 2 * cm, A4[0] - 4 * cm, 1 * cm, fill=1, stroke=0)
        canvas.setFont("Helvetica-Bold", 9)
        canvas.setFillColor(primary)
        canvas.drawString(2.2 * cm, A4[1] - 1.55 * cm, doc_type)
        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(secondary)
        canvas.drawRightString(A4[0] - 2.2 * cm, A4[1] - 1.55 * cm, version)
        # Footer
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(secondary)
        canvas.drawCentredString(
            A4[0] / 2, 1.5 * cm,
            f"{company_name} | {doc_type} | {version} | Page {doc.page}",
        )
        canvas.drawCentredString(
            A4[0] / 2, 1.1 * cm,
            "This document is confidential and intended for authorized personnel only.",
        )
        canvas.restoreState()

    story = []

    def strip_bold(text: str) -> str:
        """Convert markdown inline formatting to reportlab XML tags."""
        text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
        text = re.sub(r"\*([^*\n]+?)\*", r"<i>\1</i>", text)
        text = re.sub(r"`([^`\n]+?)`", r"<font name='Courier' size='10'>\1</font>", text)
        return text

    lines = markdown_content.splitlines()
    i = 0
    bullet_items = []
    numbered_items = []

    def flush_lists():
        nonlocal bullet_items, numbered_items
        if bullet_items:
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(strip_bold(t), bullet_style)) for t in bullet_items],
                    bulletType="bullet",
                    leftIndent=16,
                    spaceAfter=6,
                )
            )
            bullet_items = []
        if numbered_items:
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(strip_bold(t), bullet_style)) for t in numbered_items],
                    bulletType="1",
                    leftIndent=16,
                    spaceAfter=6,
                )
            )
            numbered_items = []

    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith("# "):
            flush_lists()
            story.append(Paragraph(strip_bold(line[2:].strip()), h1_style))
            story.append(HRFlowable(width="100%", thickness=1, color=accent))
            story.append(Spacer(1, 4))
            i += 1
            continue
        if line.startswith("## "):
            flush_lists()
            story.append(Spacer(1, 6))
            story.append(Paragraph(strip_bold(line[3:].strip()), h2_style))
            i += 1
            continue
        if line.startswith("### "):
            flush_lists()
            story.append(Paragraph(strip_bold(line[4:].strip()), h3_style))
            i += 1
            continue

        # Tables
        if line.startswith("|"):
            flush_lists()
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            data_rows = [
                [c.strip() for c in r.strip().strip("|").split("|")]
                for r in table_lines
                if not re.match(r"^\|[\s\-:|]+\|$", r)
            ]
            if data_rows:
                col_count = max(len(r) for r in data_rows)
                padded = [r + [""] * (col_count - len(r)) for r in data_rows]
                t = Table(padded, repeatRows=1)
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), accent),
                    ("TEXTCOLOR",  (0, 0), (-1, 0), primary),
                    ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE",   (0, 0), (-1, -1), 10),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [white, HexColor("#F8FAFC")]),
                    ("GRID",       (0, 0), (-1, -1), 0.5, HexColor("#E2E8F0")),
                    ("VALIGN",     (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]))
                story.append(t)
                story.append(Spacer(1, 8))
            continue

        # Bullets
        ul_match = re.match(r"^[-*]\s+(.*)", line)
        if ul_match:
            if numbered_items:
                flush_lists()
            bullet_items.append(ul_match.group(1))
            i += 1
            continue

        # Numbered
        ol_match = re.match(r"^\d+\.\s+(.*)", line)
        if ol_match:
            if bullet_items:
                flush_lists()
            numbered_items.append(ol_match.group(1))
            i += 1
            continue

        flush_lists()

        if not line.strip():
            story.append(Spacer(1, 4))
        elif re.match(r"^-{3,}$", line.strip()):
            story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#E2E8F0")))
        else:
            story.append(Paragraph(strip_bold(line.strip()), body_style))

        i += 1

    flush_lists()

    pdf.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
# Markdown builder (Phase 3 — trivial)
# ══════════════════════════════════════════════════════════════════════════════

def build_markdown(markdown_content: str, doc_type: str, version: str = "v1.0") -> io.BytesIO:
    """Wraps raw markdown with a metadata header and returns as UTF-8 bytes."""
    from datetime import datetime
    header = (
        f"---\n"
        f"title: {doc_type}\n"
        f"version: {version}\n"
        f"generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}\n"
        f"confidential: true\n"
        f"---\n\n"
    )
    buf = io.BytesIO()
    buf.write((header + markdown_content).encode("utf-8"))
    buf.seek(0)
    return buf
